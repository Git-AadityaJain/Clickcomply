"""
Generate DPDP-aligned privacy policy draft and export to DOCX/PDF.
"""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt

from app.schemas.org_profile import (
    AudienceType,
    DataCollected,
    DataStorageLocation,
    DpdpProcessingType,
    OrgProfile,
    RetentionChoice,
)

try:
    from fpdf import FPDF

    _HAS_FPDF = True
except ImportError:
    _HAS_FPDF = False


def _retention_label(profile: OrgProfile) -> str:
    mapping = {
        RetentionChoice.DAYS_30: "within 30 days of account deletion",
        RetentionChoice.DAYS_90: "within 90 days of account deletion",
        RetentionChoice.YEAR_1: "within one year of account deletion, unless law requires longer",
    }
    return mapping[profile.retention_period]


def generate_policy_markdown(profile: OrgProfile) -> str:
    """Build tailored privacy policy text aligned with declared DPDP obligations."""
    processing = {
        DpdpProcessingType.DIGITAL_SERVICE: "a digital service that collects personal data from users",
        DpdpProcessingType.EMPLOYEE_DATA: "an organization that processes employee and workforce personal data",
        DpdpProcessingType.VENDOR_PROCESSING: "a service that processes personal data on behalf of client organizations",
        DpdpProcessingType.MINIMAL_CONTACT: "a service that collects limited contact information",
    }[profile.processing_type]

    audience = {
        AudienceType.PUBLIC: "members of the general public",
        AudienceType.BUSINESS: "business customers and their authorized users",
        AudienceType.EMPLOYEES: "our employees and contractors",
    }[profile.audience_type]

    data = {
        DataCollected.CONTACT: "names, email addresses, and phone numbers",
        DataCollected.ACCOUNT: "account credentials, profile information, and usage tied to an account",
        DataCollected.UPLOADS: "documents and files that users choose to upload",
        DataCollected.TECHNICAL: "IP addresses, device identifiers, and technical logs",
    }[profile.data_collected]

    sections = [
        f"{profile.legal_name} — Privacy Policy & Data Protection Notice",
        "",
        f"Website: {profile.website_domain}",
        f"Privacy contact: {profile.contact_email}",
        "",
        "This draft is generated from information you provided. It is not legal advice — have a qualified lawyer review it before publishing.",
        "",
        "1. Who we are",
        f"{profile.legal_name} is the Data Fiduciary for personal data processed through {profile.website_domain}. "
        f"We operate {processing}.",
        "",
        f"Grievance officer: {profile.grievance_officer_name} ({profile.grievance_officer_designation})",
        f"Email: {profile.contact_email}",
        "",
        "2. What we do",
        f"We operate {processing} for {audience}.",
        f"The main personal data we handle is: {data}.",
        "",
        "3. Personal data we collect (DPDP Section 5; Rules 2025 Rule 3)",
        f"We collect: {data}.",
        "We collect this only for purposes needed to provide our service, keep it secure, and meet legal duties.",
        "",
        "4. Consent and lawful processing (Sections 6–7)",
        "Where the law requires consent, we ask in clear language before using personal data. "
        "You may withdraw consent as easily as you gave it. We also rely on lawful exceptions under Section 7 where applicable "
        "(for example legal compliance, fraud prevention, or employment purposes).",
        "",
    ]

    if profile.uses_ai == "yes":
        sections += [
            "5. Use of AI",
            "We use automated or AI tools in our processing. We describe the purpose, the data involved, and your rights in this notice.",
            "",
        ]

    if profile.platform_for_under_18 == "yes":
        sections += [
            "6. Children's data (Section 9)",
            "We obtain verifiable parental or guardian consent before processing children's personal data. "
            "We do not track or target advertising at children.",
            "",
        ]

    sections += [
        "7. Security (Section 8(5))",
        "We use reasonable technical and organisational measures to protect personal data against unauthorised access, "
        "loss, or misuse."
        if profile.has_security_safeguards == "yes"
        else "We are implementing safeguards appropriate to the risk of our processing.",
        "",
        "8. Retention and erasure (Section 8(7))",
        f"We keep personal data only as long as needed. After account deletion, we erase data {_retention_label(profile)}, "
        "unless a longer period is required by law.",
        "",
    ]

    if profile.data_storage_location != DataStorageLocation.INDIA or profile.users_outside_india == "yes":
        loc = {
            DataStorageLocation.INDIA: "India",
            DataStorageLocation.US: "the United States",
            DataStorageLocation.EU_UK: "the European Union or United Kingdom",
            DataStorageLocation.ASIA: "Asia-Pacific locations",
            DataStorageLocation.MULTIPLE: "multiple countries",
        }[profile.data_storage_location]
        sections += [
            "9. Cross-border processing (Section 16)",
            f"Personal data may be processed in {loc}. We follow applicable restrictions on transfers outside India.",
            "",
        ]

    if profile.uses_third_parties == "yes":
        sections += [
            "10. Third-party processors (Section 8(1)–(2))",
            "We use trusted third parties under contracts that require security and lawful processing. "
            "We remain responsible when they act on our instructions.",
            "",
        ]

    if profile.uses_analytics_cookies == "yes":
        sections += [
            "Cookies and analytics",
            "We use analytics or non-essential cookies only with appropriate notice and consent where required.",
            "",
        ]

    sections += [
        "Your rights (Sections 11–14)",
        "You may ask to access, correct, update, or erase your personal data, and to know how it is shared. "
        f"Contact us at {profile.contact_email}.",
        "",
        "Grievance redressal",
        f"Contact {profile.grievance_officer_name} at {profile.contact_email}. We aim to respond within 90 days.",
        "",
        "Complaint to the Data Protection Board",
        "If unresolved, you may complain to the Data Protection Board of India as prescribed by the Board.",
        "",
        "Data breaches",
        "If a breach affects your personal data, we will notify you and the Board without undue delay, "
        "and within 72 hours where required.",
        "",
    ]

    if profile.sells_personal_data or profile.shares_for_advertising:
        sections += [
            "Important",
            "Our activities include selling or sharing personal data for advertising — this requires explicit disclosure and lawful basis.",
            "",
        ]

    return "\n".join(sections)


def export_policy_docx(markdown: str, title: str) -> bytes:
    doc = DocxDocument()
    doc.add_heading(title, level=0)
    for block in markdown.split("\n\n"):
        for line in block.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line[0].isdigit() and ". " in line[:4]:
                doc.add_heading(line, level=2)
            else:
                p = doc.add_paragraph(line)
                p.style.font.size = Pt(11)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_policy_pdf(markdown: str, title: str) -> bytes:
    if not _HAS_FPDF:
        raise RuntimeError("PDF export is not available on this server.")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.multi_cell(0, 8, title)
    pdf.ln(4)
    pdf.set_font("Helvetica", size=10)
    for line in markdown.split("\n"):
        text = line.strip()
        if not text:
            pdf.ln(3)
            continue
        pdf.multi_cell(0, 5, text)
    out = pdf.output()
    return out if isinstance(out, bytes) else out.encode("latin-1")
