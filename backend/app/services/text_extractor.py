"""
Extract plain text from uploaded compliance documents (PDF and DOCX).
"""

from io import BytesIO
from pathlib import Path

from app.core.logging import logger


class TextExtractionError(Exception):
    """Raised when document text cannot be extracted."""


def extract_text_from_bytes(file_content: bytes, filename: str) -> str:
    """
    Extract text from file bytes based on extension.

    Args:
        file_content: Raw file bytes.
        filename: Original filename (used to detect format).

    Returns:
        Extracted plain text.

    Raises:
        TextExtractionError: Unsupported format or extraction failure.
    """
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _extract_pdf(file_content)
    if lower.endswith(".docx"):
        return _extract_docx(file_content)
    if lower.endswith(".doc"):
        raise TextExtractionError(
            "Legacy .doc format is not supported. Please upload PDF or DOCX."
        )
    raise TextExtractionError(f"Unsupported file type: {filename}")


def extract_text_from_path(file_path: Path) -> str:
    """Extract text from a file on disk."""
    content = file_path.read_bytes()
    return extract_text_from_bytes(content, file_path.name)


def _extract_pdf(file_content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise TextExtractionError("pypdf is not installed") from exc

    try:
        reader = PdfReader(BytesIO(file_content))
        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text.strip())
        combined = "\n\n".join(p for p in pages if p)
        if not combined.strip():
            raise TextExtractionError("PDF contains no extractable text")
        return combined
    except TextExtractionError:
        raise
    except Exception as exc:
        logger.error(f"PDF extraction failed: {exc}")
        raise TextExtractionError(f"Failed to read PDF: {exc}") from exc


def _extract_docx(file_content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise TextExtractionError("python-docx is not installed") from exc

    try:
        doc = Document(BytesIO(file_content))
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style else ""
            # Preserve heading structure so LLM can locate sections by name
            if style.startswith("Heading 1") or style.startswith("Title"):
                parts.append(f"\n## {text}")
            elif style.startswith("Heading 2"):
                parts.append(f"\n### {text}")
            elif style.startswith("Heading"):
                parts.append(f"\n#### {text}")
            else:
                parts.append(text)

        # Extract table content (DPO tables, contact tables, etc.)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        combined = "\n\n".join(parts)
        if not combined.strip():
            raise TextExtractionError("DOCX contains no extractable text")
        return combined
    except TextExtractionError:
        raise
    except Exception as exc:
        logger.error(f"DOCX extraction failed: {exc}")
        raise TextExtractionError(f"Failed to read DOCX: {exc}") from exc
