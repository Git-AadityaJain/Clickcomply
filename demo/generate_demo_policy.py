"""
Generate a realistic but intentionally incomplete demo privacy policy DOCX.

This policy is written in the same professional style and structure as the
perfect policy, but has deliberate gaps that the AI should flag:

  EXPECTED FAILURES:
    RULE_GRIEVANCE_001  FAIL          -- no grievance officer section or 90-day timeline
    RULE_BREACH_001     FAIL          -- no data breach notification section
    RULE_BOARD_001      FAIL          -- no Board complaint information
    RULE_RETENTION_001  NEEDS_REVIEW  -- says "reasonable period", not a specific timeframe
    RULE_RIGHTS_001     NEEDS_REVIEW  -- only access + correction; erasure and nomination absent

  EXPECTED PASSES:
    RULE_NOTICE_001     PASS
    RULE_CONSENT_001    PASS
    RULE_LEGITIMATE_001 PASS
    RULE_SECURITY_001   PASS
    RULE_VENDOR_001     PASS
    RULE_TRANSFER_001   PASS
    RULE_SDF_001        PASS
    RULE_CONTACT_001    NEEDS_REVIEW  -- email provided but not formally designated per Rule 9

  NOT APPLICABLE (with these questionnaire answers):
    RULE_CHILD_001      -- Under 18: No
    RULE_DISABILITY_001 -- Under 18: No
    RULE_CONSENT_MGR_001-- Uses AI: No

Questionnaire answers to use when uploading:
  Company:          NovaStack Solutions Pvt Ltd
  Website:          https://novastack.example
  Email:            privacy@novastack.example
  Grievance:        Priya Sharma
  Processing:       Website / app with user accounts
  Audience:         Business customers
  Data:             Account and login details
  Uses AI:          No
  Data storage:     India
  Third parties:    Yes
  Analytics:        No
  Retention:        90 days after account closure
  Under 18:         No
  Intl users:       No
  Security:         Yes
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

TITLE = "NovaStack Solutions Pvt Ltd - Privacy Policy"
META  = "Effective Date: 1 March 2025  |  Version: 1.0"
OUTPUT = Path(__file__).resolve().parent / "NovaStack_Demo_Privacy_Policy.docx"

BODY = [
    # Section 1 — Overview
    (1, "1. Overview and Scope"),
    (None,
     "NovaStack Solutions Pvt Ltd ('NovaStack', 'we', 'us') operates a B2B project "
     "management platform accessible at https://novastack.example. This Privacy Policy "
     "describes what personal data we collect from our business customers, why we collect "
     "it, and how we protect it under India's Digital Personal Data Protection Act, 2023 "
     "('DPDP Act') and the DPDP Rules 2025.\n\n"
     "This policy applies to all personal data processed by NovaStack in connection with "
     "our platform and associated services."),

    # Section 2 — Notice (PASS: itemised data, purposes, withdrawal method described)
    (1, "2. Personal Data We Collect and Why (Notice - Section 5 DPDP Act)"),
    (None,
     "We collect the following categories of personal data when you register or use our platform:\n\n"
     "  a) Account and login credentials: full name, work email address, company name, "
     "username, and password (stored as a salted hash) — to create your account, "
     "authenticate you, and deliver the project management service.\n\n"
     "  b) Profile and billing details: job title, business address, and payment reference "
     "identifiers — to manage your subscription and issue invoices.\n\n"
     "  c) Usage and activity logs: actions taken within the platform such as tasks created, "
     "comments posted, and files shared — to provide the service and enable team "
     "collaboration.\n\n"
     "We do not sell your personal data. We do not share your data for advertising or "
     "behavioural targeting.\n\n"
     "How to withdraw consent: You may withdraw consent at any time by deleting your account "
     "via Settings or by writing to privacy@novastack.example. Upon withdrawal, we will "
     "cease non-essential processing within thirty (30) days.\n\n"
     "How to exercise your rights: See Section 6 of this policy."),

    # Section 3 — Consent (PASS: affirmative consent, withdrawal, 30-day cessation)
    (1, "3. Legal Basis for Processing - Consent (Section 6 DPDP Act)"),
    (None,
     "We process your personal data based on your consent. When you register, you give "
     "free, specific, informed, and unambiguous consent by affirmatively accepting this "
     "Privacy Policy through a clear checkbox action.\n\n"
     "Withdrawal of consent: You may withdraw your consent at any time by deleting your "
     "account or writing to privacy@novastack.example. Withdrawal does not affect the "
     "lawfulness of processing carried out before withdrawal.\n\n"
     "Cessation within 30 days: Upon receiving a valid consent withdrawal request, "
     "NovaStack will cease all non-essential processing of your personal data within "
     "thirty (30) days. Processing required to fulfil a legal obligation or resolve an "
     "existing dispute may continue only to the extent strictly necessary.\n\n"
     "We do not use Consent Managers; all consent is obtained and managed directly "
     "through our platform."),

    # Section 4 — Legitimate uses (PASS)
    (1, "4. Legal Basis for Processing - Legitimate Uses Without Consent (Section 7 DPDP Act)"),
    (None,
     "We may process personal data without consent to the extent required to:\n\n"
     "  a) Comply with a court order, legal obligation, or direction of a competent "
     "authority (for example, tax or audit requirements).\n\n"
     "  b) Respond to a medical or public safety emergency affecting a user.\n\n"
     "  c) Fulfil obligations in connection with a merger, acquisition, or asset transfer, "
     "subject to equivalent data protection safeguards.\n\n"
     "We will document the basis for such processing and inform affected users where "
     "disclosure is not restricted by law."),

    # Section 5 — Vendors (PASS: processors named, DPAs mentioned, fiduciary responsibility)
    (1, "5. Third-Party Data Processors (Section 8 DPDP Act)"),
    (None,
     "To deliver our platform, we engage the following categories of third-party Data "
     "Processors, each operating only under our written instructions and under a Data "
     "Processing Agreement that imposes security, confidentiality, and breach notification "
     "obligations:\n\n"
     "  - Cloud infrastructure provider: hosts the platform and databases on servers "
     "located exclusively in India.\n\n"
     "  - Transactional email provider: delivers account verification and service "
     "notifications. Processes only the recipient email address and message content.\n\n"
     "  - Payment gateway: processes subscription billing under PCI-DSS controls. We do "
     "not store full card numbers.\n\n"
     "NovaStack remains fully responsible as the Data Fiduciary for all processing by "
     "these processors."),

    # Section 6 — Transfer (PASS: India only, no cross-border)
    (1, "6. Data Storage Location and Cross-Border Transfers (Section 16 DPDP Act)"),
    (None,
     "All personal data is stored and processed exclusively on servers located in India. "
     "We do not transfer personal data outside India. We do not send personal data to any "
     "country restricted by the Central Government under Section 16 of the DPDP Act. "
     "Should we need to make cross-border transfers in future, we will update this policy "
     "and obtain any required approvals before doing so."),

    # Section 7 — Security (PASS: specific measures listed)
    (1, "7. Security Safeguards (Section 8(5) DPDP Act and DPDP Rules 2025 Rule 6)"),
    (None,
     "We implement the following technical and organisational security safeguards to "
     "protect personal data:\n\n"
     "  - Encryption in transit: all data transmitted between your browser and our servers "
     "is encrypted using TLS 1.2 or higher.\n\n"
     "  - Encryption at rest: personal data is encrypted using AES-256.\n\n"
     "  - Access controls: role-based access controls (RBAC) ensure only authorised "
     "personnel access personal data, based on the principle of least privilege. Access "
     "rights are reviewed quarterly.\n\n"
     "  - Monitoring and logging: we maintain audit logs of access to personal data "
     "systems, retained for a minimum of one year.\n\n"
     "  - Organisational measures: all staff with data access complete annual data "
     "protection training and are bound by confidentiality obligations."),

    # Section 8 — Retention (NEEDS_REVIEW: "reasonable period" is vague — deliberate gap)
    (1, "8. Retention and Erasure (Section 8(7)-(8) DPDP Act and DPDP Rules 2025 Rule 8)"),
    (None,
     "We retain personal data only for as long as necessary for the stated purposes or as "
     "required by law:\n\n"
     "  - Active accounts: personal data is retained while your account is active and "
     "your subscription is in force.\n\n"
     "  - After account closure: when you delete your account, we will erase your personal "
     "data within a reasonable period, unless retention is required by applicable law.\n\n"
     "  - Inactive accounts: if your account has been inactive for an extended period, we "
     "may contact you before deleting your data.\n\n"
     "  - System logs: audit logs are retained for a minimum of one year.\n\n"
     "When data is erased, it is deleted from primary databases and purged from backups "
     "within the next scheduled backup cycle."),

    # Section 9 — Rights (NEEDS_REVIEW: only access + correction; no erasure or nomination — deliberate gap)
    (1, "9. Your Rights as a Data Principal (Sections 11-14 DPDP Act and DPDP Rules 2025 Rule 14)"),
    (None,
     "You have the following rights regarding your personal data. To exercise any right, "
     "write to privacy@novastack.example with the subject line 'Data Rights Request':\n\n"
     "  a) Right to access: request a summary of the personal data we hold about you.\n\n"
     "  b) Right to correction: request correction or updating of any inaccurate or "
     "incomplete personal data we hold about you.\n\n"
     "We will acknowledge your request within 7 days and respond within 30 days. We do "
     "not charge a fee for reasonable requests."),

    # Section 10 — SDF (PASS: explicitly not an SDF)
    (1, "10. Significant Data Fiduciary Status (Section 10 DPDP Act)"),
    (None,
     "NovaStack Solutions Pvt Ltd has NOT been designated as a Significant Data Fiduciary "
     "(SDF) by the Central Government under Section 10 of the DPDP Act. SDF obligations "
     "such as mandatory DPO appointment, periodic DPIAs, and independent audits do not "
     "currently apply to NovaStack. We will monitor any change in our SDF designation and "
     "implement required obligations promptly if designated in future."),

    # Section 11 — Contact (NEEDS_REVIEW: email provided but no formal Rule 9 designation — deliberate gap)
    (1, "11. Contact Us"),
    (None,
     "For any questions about this Privacy Policy or about how we handle your personal "
     "data, please email us at privacy@novastack.example. We aim to respond to all "
     "enquiries within 5 business days."),

    # Section 12 — Updates
    (1, "12. Updates to This Policy"),
    (None,
     "This policy may be updated from time to time. Material changes will be communicated "
     "to registered users by email before they take effect. Continued use of the platform "
     "after a policy update constitutes acceptance of the revised terms. The current "
     "version is always available at https://novastack.example/privacy."),

    # -----------------------------------------------------------------------
    # DELIBERATELY OMITTED sections (these create the AI-detected failures):
    #   - Grievance Redressal section   → RULE_GRIEVANCE_001 FAIL
    #   - Data Breach Notification      → RULE_BREACH_001    FAIL
    #   - Board Complaint section       → RULE_BOARD_001     FAIL
    # -----------------------------------------------------------------------
]


def main() -> None:
    doc = Document()

    heading = doc.add_heading(TITLE, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(META)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.size = Pt(9)
        run.font.italic = True

    for level, text in BODY:
        if level == 1:
            doc.add_heading(text, level=1)
        else:
            doc.add_paragraph(text)

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print()
    print("=== QUESTIONNAIRE ANSWERS TO USE ===")
    print("Company:       NovaStack Solutions Pvt Ltd")
    print("Website:       https://novastack.example")
    print("Email:         privacy@novastack.example")
    print("Grievance:     Priya Sharma")
    print("Processing:    Website / app with user accounts")
    print("Audience:      Business customers")
    print("Data:          Account and login details")
    print("Uses AI:       No")
    print("Storage:       India")
    print("Third parties: Yes")
    print("Analytics:     No")
    print("Retention:     90 days after account closure")
    print("Under 18:      No")
    print("Intl users:    No")
    print("Security:      Yes")
    print()
    print("=== EXPECTED ANALYSIS OUTPUT ===")
    print("FAIL:         Grievance Redressal, Breach Notification, Board Complaint")
    print("NEEDS_REVIEW: Retention (vague period), Rights (incomplete), Contact (informal)")
    print("PASS:         Notice, Consent, Legitimate Uses, Security, Processors, Transfer, SDF")
    print("N/A:          Children, Disability, Consent Manager")


if __name__ == "__main__":
    main()
