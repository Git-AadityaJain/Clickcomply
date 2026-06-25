"""
Generate a DPDP-compliant privacy policy DOCX for NovaStack Solutions.
This policy is designed to PASS all 14 applicable compliance rules when
uploaded against the standard NovaStack demo questionnaire answers.

Questionnaire answers used:
  Company:        NovaStack Solutions Pvt Ltd
  Website:        https://novastack.example
  Email:          privacy@novastack.example
  Grievance:      Priya Sharma (Grievance Officer)
  Processing:     Website / app with user accounts
  Audience:       Business customers
  Data:           Login and account details
  Storage:        India
  Uses AI:        No
  Third parties:  Yes
  Analytics:      No
  Retention:      90 days after account closure
  Under 18:       No
  International:  No
  Security:       Yes
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = Path(__file__).resolve().parent / "NovaStack_Perfect_Privacy_Policy.docx"

TITLE = "NovaStack Solutions Pvt Ltd - Privacy Policy"
META  = "Effective Date: 1 January 2025  |  Version: 2.0"

BODY = [
    (1, "1. Overview and Scope"),
    (None,
     "NovaStack Solutions Pvt Ltd (\"NovaStack\", \"we\", \"us\") operates a B2B project "
     "management platform accessible at https://novastack.example. This Privacy Policy "
     "describes, in plain language, what personal data we collect from our business "
     "customers, why we collect it, how we protect it, and your rights under India's "
     "Digital Personal Data Protection Act, 2023 (\"DPDP Act\") and the DPDP Rules 2025.\n\n"
     "This policy is available in English. Where we are required to provide it in any "
     "language listed in the Eighth Schedule of the Constitution of India, we will do so "
     "upon request to privacy@novastack.example."),

    (1, "2. Personal Data We Collect and Why (Notice - Section 5 DPDP Act)"),
    (None,
     "We collect the following categories of personal data when you register or use our platform:\n\n"
     "  a) Account and login credentials - full name, work email address, company name, "
     "username, and password (stored as a salted hash) - to create your account, authenticate "
     "you, and deliver the project management service.\n\n"
     "  b) Profile and billing details - job title, business address, and payment reference "
     "identifiers - to manage your subscription and issue invoices.\n\n"
     "  c) Usage and activity logs - actions taken within the platform such as tasks created, "
     "comments posted, and files shared - to provide the service and enable team collaboration.\n\n"
     "We do not collect sensitive personal data, financial account numbers, health records, "
     "or biometric identifiers. We do not sell your personal data. We do not share your "
     "personal data for advertising or behavioural targeting.\n\n"
     "You may withdraw consent at any time by deleting your account (see Section 8). You may "
     "also exercise your rights or raise concerns by contacting privacy@novastack.example. "
     "If your concerns are not resolved, you may complain to the Data Protection Board of "
     "India (see Section 12)."),

    (1, "3. Legal Basis for Processing (Consent and Legitimate Uses - Sections 6-7)"),
    (None,
     "We process your personal data on the following legal bases:\n\n"
     "  a) Consent (Section 6): When you register, you give free, specific, informed, and "
     "unambiguous consent by affirmatively accepting this policy and our Terms of Service. "
     "Consent is obtained through a clear affirmative act and is not bundled with other terms "
     "in a way that prevents separate refusal.\n\n"
     "Withdrawal of consent: You may withdraw your consent at any time by deleting your "
     "account via the platform settings or by sending a written withdrawal request to "
     "privacy@novastack.example. Withdrawal of consent will not affect the lawfulness of "
     "any processing carried out before the withdrawal.\n\n"
     "Cessation of processing within 30 days: Upon receiving a valid consent withdrawal "
     "request, NovaStack will cease all non-essential processing of your personal data within "
     "thirty (30) days. Processing that is required to fulfil a legal obligation, resolve an "
     "existing dispute, or enforce our agreements may continue beyond this period only to the "
     "extent strictly necessary and permitted by law.\n\n"
     "What happens to your data after withdrawal: Once non-essential processing ceases, your "
     "personal data will be erased within 90 days of account closure, subject to any mandatory "
     "legal retention obligations (for example, financial records under the Companies Act). "
     "We will send you a confirmation email once your data has been deleted. During the 30-day "
     "cessation period your data will not be used for any new processing activities; it will "
     "only be retained in read-only form pending deletion.\n\n"
     "  b) Legitimate uses without consent (Section 7): We may process personal data without "
     "consent to the extent required to comply with a court order, a legal obligation, or to "
     "respond to a medical or public safety emergency. We will document the basis and inform "
     "you where disclosure is not restricted by law.\n\n"
     "We do not use Consent Managers; all consent is obtained and managed directly through "
     "our platform. We do not use automated decision-making that produces legal or similarly "
     "significant effects without human review."),

    (1, "4. Third-Party Data Processors (Section 8 DPDP Act)"),
    (None,
     "To deliver the platform, we engage the following categories of third-party Data Processors. "
     "Each processes personal data only under our written instructions and under a Data Processing "
     "Agreement that imposes security, confidentiality, sub-processor restrictions, audit rights, "
     "and breach notification obligations:\n\n"
     "  - Cloud infrastructure provider: hosts the platform and its databases on servers located "
     "exclusively in India (Mumbai region). They may not process data for their own purposes.\n\n"
     "  - Transactional email provider: delivers account verification, password reset, and service "
     "notifications. They process only the recipient email address and message content necessary "
     "for delivery.\n\n"
     "  - Payment gateway: processes subscription billing. We do not store full card numbers; "
     "the gateway handles payment data under PCI-DSS controls.\n\n"
     "NovaStack remains fully responsible as the Data Fiduciary for all processing by these "
     "processors. We periodically review processor compliance and will not engage any processor "
     "that cannot provide sufficient data protection guarantees. Each processor is contractually "
     "required to apply security safeguards at least equivalent to those we apply directly."),

    (1, "5. Data Storage Location and Cross-Border Transfers (Section 16 DPDP Act)"),
    (None,
     "All personal data collected through our platform is stored and processed exclusively on "
     "servers located in India. We do not transfer personal data outside India. We do not send "
     "personal data to any country restricted by the Central Government under Section 16 of the "
     "DPDP Act. Should we need to transfer data internationally in future, we will update this "
     "policy, obtain any required government approvals, and inform affected users before any "
     "such transfer takes place."),

    (1, "6. Security Safeguards (Section 8(5) DPDP Act and DPDP Rules 2025 Rule 6)"),
    (None,
     "We implement reasonable and appropriate technical and organisational security measures "
     "to protect personal data from unauthorised access, disclosure, alteration, or destruction:\n\n"
     "  - Encryption in transit: all data transmitted between your browser or client and our "
     "servers is encrypted using TLS 1.2 or higher.\n\n"
     "  - Encryption at rest: personal data stored in our databases is encrypted using AES-256 "
     "or equivalent industry-standard encryption.\n\n"
     "  - Access controls: role-based access controls (RBAC) ensure only authorised personnel "
     "access personal data, based on the principle of least privilege. Access privileges are "
     "reviewed quarterly.\n\n"
     "  - Monitoring and logging: we maintain audit logs of access to personal data systems. "
     "Logs are retained for a minimum of one year and monitored continuously for signs of "
     "unauthorised access or anomalies.\n\n"
     "  - Backup and business continuity: we maintain encrypted off-site backups and a tested "
     "business continuity plan to ensure data availability and rapid recovery.\n\n"
     "  - Organisational measures: all employees with access to personal data complete mandatory "
     "data protection training annually and are bound by confidentiality obligations. Third-party "
     "access is governed by Data Processing Agreements.\n\n"
     "  - Data accuracy: we take reasonable steps to ensure personal data used in decisions "
     "that affect you is accurate and kept up to date. You may correct inaccurate data at "
     "any time by writing to privacy@novastack.example."),

    (1, "7. Retention and Erasure (Section 8(7)-(8) DPDP Act and DPDP Rules 2025 Rule 8)"),
    (None,
     "We retain personal data only as long as necessary for the stated purposes or as required by law:\n\n"
     "  - Active accounts: personal data is retained while your account is active and your "
     "subscription is in force.\n\n"
     "  - After account closure: when you delete your account or your subscription ends, we "
     "will erase your personal data within 90 days, unless retention is required by law "
     "(for example, financial records required under the Companies Act or GST law).\n\n"
     "  - Inactive accounts: if your account has been inactive (no login or activity) for "
     "24 continuous months, we will send a notice to your registered email address at least "
     "30 days before deleting your data, giving you the opportunity to log in to reactivate "
     "your account and prevent deletion.\n\n"
     "  - Processing and access logs: system audit logs are retained for a minimum of one "
     "year from creation as required under DPDP Rules 2025 Rule 8, after which they are "
     "securely and irreversibly deleted.\n\n"
     "  - Consent records: records evidencing your consent are retained for the duration of "
     "our relationship and for three years thereafter in case of legal dispute.\n\n"
     "When data is erased, it is deleted from primary databases and purged from backups "
     "within the next scheduled backup cycle (maximum 30 additional days)."),

    (1, "8. Your Rights as a Data Principal (Sections 11-14 DPDP Act and DPDP Rules 2025 Rule 14)"),
    (None,
     "You have the following rights regarding your personal data. To exercise any right, "
     "write to privacy@novastack.example with the subject line 'Data Rights Request' "
     "and include your name and registered email address:\n\n"
     "  a) Right to access: request a summary of the personal data we hold about you and "
     "a list of Data Processors and third parties with whom it has been shared.\n\n"
     "  b) Right to correction: request correction, completion, or updating of any inaccurate, "
     "incomplete, or outdated personal data we hold about you.\n\n"
     "  c) Right to erasure: request deletion of your personal data where it is no longer "
     "necessary for the purposes collected, subject to legal retention requirements.\n\n"
     "  d) Right to withdraw consent: withdraw consent at any time by deleting your account "
     "or sending a written withdrawal request to privacy@novastack.example. Withdrawal does "
     "not affect the lawfulness of processing carried out before withdrawal.\n\n"
     "  e) Right to information about processors: request details of the Data Processors to "
     "whom your personal data has been disclosed and for what purpose.\n\n"
     "  f) Right of nomination: nominate another individual to exercise your data rights on "
     "your behalf in the event of your death or incapacity, by submitting a written nomination "
     "to privacy@novastack.example.\n\n"
     "We will acknowledge your request within 7 days and respond fully within 30 days. "
     "We do not charge a fee for reasonable requests. We may ask you to verify your identity "
     "before processing your request."),

    (1, "9. Data of Persons with Disability - Guardian Authority Verification (Section 9(1) DPDP Act and DPDP Rules 2025 Rule 11)"),
    (None,
     "Our platform is designed for business customers and is not directed at individuals under 18. "
     "Where a user is a person with disability who has a lawful guardian, NovaStack will process "
     "their personal data only after obtaining verifiable consent from that guardian and completing "
     "a formal guardian authority verification process, as required by Section 9(1) of the DPDP "
     "Act and Rule 11 of the DPDP Rules 2025.\n\n"
     "Guardian Authority Verification Process - Documents Required and Criteria:\n\n"
     "Before accepting any guardian's consent, we require the following documents as evidence of "
     "lawful guardianship authority:\n"
     "  (a) Proof of identity: a valid government-issued photo identity document for both the "
     "guardian (e.g., Aadhaar, PAN, passport) and the data principal.\n"
     "  (b) Legal authority document: a court order, notarised guardianship deed, power of "
     "attorney, or any other legal instrument issued by a competent authority establishing the "
     "guardian's legal authority over the data principal.\n"
     "  (c) Disability certificate: a valid disability certificate or equivalent document issued "
     "by a competent medical or government authority confirming the nature of the disability.\n\n"
     "Verification criteria: we will assess whether (i) the legal instrument is valid, current, "
     "and not revoked; (ii) the identity of the guardian matches the instrument; and (iii) the "
     "scope of authority covers consent to personal data processing.\n\n"
     "Verification procedure:\n"
     "  Step 1 - Request: the guardian emails privacy@novastack.example with subject 'Guardian "
     "Consent Verification', providing the data principal's name, disability details, and their "
     "own contact information.\n"
     "  Step 2 - Document submission: guardian submits the required documents listed above.\n"
     "  Step 3 - Review: NovaStack reviews submissions within 10 business days against the "
     "verification criteria and confirms in writing whether authority is verified.\n"
     "  Step 4 - Processing: personal data is processed under guardian consent only after "
     "written confirmation of successful verification.\n"
     "  Step 5 - Records: verified records are retained securely; the guardian may revoke "
     "authority at any time by written notice to privacy@novastack.example.\n\n"
     "We will not process personal data of a person with disability under guardian consent until "
     "all verification steps are completed. If verification fails, we will notify the guardian "
     "and will not proceed."),

    (1, "10. Data Breach Notification (Section 8(6) DPDP Act and DPDP Rules 2025 Rule 7)"),
    (None,
     "Board Intimation within 72 Hours: NovaStack will notify the Data Protection Board of India "
     "within seventy-two (72) hours of becoming aware of any personal data breach, as required by "
     "Section 8(6) of the DPDP Act and Rule 7 of the DPDP Rules 2025. This 72-hour Board "
     "intimation obligation applies regardless of whether the breach affects a small or large "
     "number of Data Principals.\n\n"
     "In the event of a personal data breach, NovaStack will take the following steps:\n\n"
     "  a) Notify affected Data Principals without undue delay: we will inform each affected "
     "individual of the nature and description of the breach, the categories and approximate "
     "volume of personal data affected, the likely consequences of the breach, the measures we "
     "have taken or propose to take to address and mitigate the breach, and a contact point for "
     "further information and assistance.\n\n"
     "  b) Notify the Data Protection Board of India within 72 hours: within seventy-two (72) "
     "hours of becoming aware of a breach, we will submit a formal breach notification to the "
     "Data Protection Board of India in the prescribed format under DPDP Rules 2025 Rule 7, "
     "including a description of the breach, its likely consequences, the data and persons "
     "affected, and the remedial measures taken or planned. If all required details are not "
     "available within 72 hours, we will submit an initial notification within that window and "
     "provide a complete follow-up report as soon as practicable.\n\n"
     "  c) Contain and investigate: we will immediately isolate affected systems, conduct a "
     "root-cause investigation, and implement remedial measures to prevent recurrence.\n\n"
     "  d) Maintain breach records: we maintain an internal register of all personal data "
     "breaches, including those that did not require external notification, for audit and "
     "accountability purposes."),

    (1, "11. Grievance Redressal (Section 8(10) and Section 13 DPDP Act and DPDP Rules 2025 Rule 14(3))"),
    (None,
     "If you have any complaint or concern about how we handle your personal data, please contact "
     "our Grievance Officer:\n\n"
     "  Name:         Priya Sharma\n"
     "  Designation:  Grievance Officer\n"
     "  Organisation: NovaStack Solutions Pvt Ltd\n"
     "  Email:        privacy@novastack.example\n"
     "  Website:      https://novastack.example/privacy\n\n"
     "We will acknowledge your grievance within 7 days and endeavour to resolve it within ninety "
     "(90) days of receipt. If your grievance is not resolved to your satisfaction within ninety "
     "days, you may escalate your complaint to the Data Protection Board of India (see Section 12). "
     "Data Principals are required to exhaust our grievance redressal process before approaching "
     "the Board, unless the Board directs otherwise in a particular case."),

    (1, "12. Complaint to the Data Protection Board of India (Section 5 DPDP Act)"),
    (None,
     "If your grievance with NovaStack is not resolved within ninety days, or if you believe we "
     "have violated the Digital Personal Data Protection Act, 2023, you have the right to file a "
     "complaint with the Data Protection Board of India.\n\n"
     "The Data Protection Board of India is the independent adjudicatory authority established "
     "under Chapter VI of the DPDP Act. Complaints may be filed through the Board's online portal "
     "or in such other manner as may be prescribed. For current guidance on filing a complaint, "
     "refer to the Board's official website or contact the Ministry of Electronics and Information "
     "Technology (MeitY) at www.meity.gov.in. Filing a complaint with the Board is free of charge."),

    (1, "13. Significant Data Fiduciary Status (Section 10 DPDP Act)"),
    (None,
     "SDF Obligations Do NOT Apply to NovaStack: NovaStack Solutions Pvt Ltd confirms that it "
     "has NOT been designated as a Significant Data Fiduciary (SDF) by the Central Government "
     "of India under Section 10 of the Digital Personal Data Protection Act, 2023. "
     "Significant Data Fiduciary obligations do not apply to NovaStack at this time.\n\n"
     "Specifically, the following SDF obligations under Section 10 and DPDP Rules 2025 Rule 13 "
     "do NOT currently apply to NovaStack Solutions Pvt Ltd:\n"
     "  - Mandatory appointment of an India-resident Data Protection Officer (DPO);\n"
     "  - Periodic Data Protection Impact Assessments (DPIA);\n"
     "  - Independent audits of data processing activities;\n"
     "  - Algorithmic or technical measures to prevent risk to Data Principal rights;\n"
     "  - Restrictions on cross-border transfer of notified sensitive personal data.\n\n"
     "We will monitor any change in our designation by the Central Government. If NovaStack is "
     "designated as a Significant Data Fiduciary in future, we will promptly implement all "
     "required SDF obligations and update this policy accordingly."),

    (1, "14. Designated Contact for Personal Data Processing Questions (Section 8(9) DPDP Act and DPDP Rules 2025 Rule 9)"),
    (None,
     "As required by Section 8(9) of the DPDP Act and Rule 9 of the DPDP Rules 2025, we have "
     "designated the following person to answer any questions about the processing of personal "
     "data by NovaStack Solutions Pvt Ltd. This contact is prominently published on our website, "
     "application, and in this policy:\n\n"
     "  Designated Contact for Data Processing Questions:\n"
     "  Name:         Priya Sharma\n"
     "  Designation:  Data Protection Contact / Grievance Officer\n"
     "  Organisation: NovaStack Solutions Pvt Ltd\n"
     "  Email:        privacy@novastack.example\n"
     "  Website:      https://novastack.example/privacy\n"
     "  Available:    Monday to Friday, 9:00 AM to 6:00 PM IST\n\n"
     "You may contact this person for any question relating to: what personal data we collect; "
     "why and how we process it; who we share it with; how long we keep it; how to exercise "
     "your rights; or any other matter concerning personal data processing by NovaStack."),

    (1, "15. Updates to This Policy"),
    (None,
     "This policy is reviewed at least annually. Material changes will be communicated to "
     "registered users by email at least 30 days before they take effect. Continued use of "
     "the platform after the effective date of a revised policy constitutes acceptance of the "
     "revised terms. The current version is always available at https://novastack.example/privacy."),
]


def main() -> None:
    doc = Document()

    heading = doc.add_heading(TITLE, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(META)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.size = Pt(9)
        run.font.italic = True

    for level, text in BODY:
        if level == 1:
            doc.add_heading(text, level=1)
        else:
            doc.add_paragraph(text)

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print("Upload this file with the standard NovaStack questionnaire answers.")
    print("Expected result: COMPLIANT or NEEDS_REVIEW with minimal/no FAIL gaps.")


if __name__ == "__main__":
    main()
