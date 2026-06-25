"""
Generate a DPDP-compliant privacy policy DOCX for GlobalLearn India Pvt Ltd.

This policy is designed to PASS all 16 compliance rules — every rule is
APPLICABLE (none are skipped). Upload it with the questionnaire answers below.

Questionnaire answers to use:
  Company:          GlobalLearn India Pvt Ltd
  Website:          https://globallearn.example
  Email:            privacy@globallearn.example
  Grievance:        Arjun Mehta (Grievance Officer / DPO)
  Processing:       Website / app with user accounts
  Audience:         General public                  <-- makes CHILD + DISABILITY rules apply
  Data:             Account and login details
  Uses AI:          Yes                             <-- makes CONSENT_MANAGER rule apply
  Data storage:     Multiple countries              <-- makes TRANSFER rule REQUIRED
  Third parties:    Yes                             <-- makes VENDOR rule REQUIRED
  Analytics:        Yes
  Retention:        90 days after account closure
  Under 18:         Yes                             <-- makes CHILD_001 REQUIRED
  Intl users:       Yes                             <-- reinforces TRANSFER rule
  Security:         Yes
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = Path(__file__).resolve().parent / "GlobalLearn_Universal_Privacy_Policy.docx"

TITLE = "GlobalLearn India Pvt Ltd - Privacy Policy"
META  = "Effective Date: 1 January 2025  |  Version: 3.0"

BODY = [
    # Section 1 — Overview (heading keywords: overview, scope)
    (1, "1. Overview and Scope"),
    (None,
     "GlobalLearn India Pvt Ltd ('GlobalLearn', 'we', 'us') operates an AI-powered online "
     "learning platform accessible at https://globallearn.example. Our platform serves learners "
     "of all ages, including children under 18, across India and internationally. This Privacy "
     "Policy describes, in plain language, what personal data we collect, why we collect it, "
     "how we protect it, and your rights under India's Digital Personal Data Protection Act, "
     "2023 ('DPDP Act') and the DPDP Rules 2025.\n\n"
     "This policy applies to all personal data processed by GlobalLearn in the course of "
     "operating our platform, mobile applications, and associated services. This policy is "
     "available in English. Translations in Eighth Schedule languages are available on request "
     "at privacy@globallearn.example."),

    # Section 2 — Notice (heading keywords: data we collect, notice, section 5)
    (1, "2. Personal Data We Collect and Why (Notice - Section 5 DPDP Act)"),
    (None,
     "We collect the following categories of personal data when you register or use our platform:\n\n"
     "  a) Account and identity details: full name, date of birth, work or personal email address, "
     "username, and password (stored as a salted cryptographic hash). We use this data to create "
     "your account, verify your identity, and deliver the learning service.\n\n"
     "  b) Profile and educational data: learning preferences, course enrolments, assessment "
     "scores, certificates earned, and language preferences. We use this data to personalise "
     "your learning experience and issue recognised credentials.\n\n"
     "  c) Contact and billing details: phone number, billing address, and payment reference "
     "identifiers. We use this data to manage subscriptions, issue invoices, and provide "
     "customer support.\n\n"
     "  d) Technical and usage data: IP address, device type, browser, pages visited, time "
     "spent on lessons, and interaction logs. We use this data to operate and improve the "
     "platform, detect security threats, and generate anonymised analytics.\n\n"
     "  e) AI personalisation data: learner behaviour signals (quiz results, replay patterns, "
     "completion rates) used by our AI engine to recommend courses and adapt content difficulty.\n\n"
     "We do not sell your personal data. We do not share your data for third-party advertising "
     "or behavioural targeting unrelated to our educational services.\n\n"
     "How to withdraw consent: You may withdraw consent at any time by deleting your account "
     "via Settings or by writing to privacy@globallearn.example. Withdrawal does not affect "
     "the lawfulness of any processing carried out before withdrawal.\n\n"
     "How to exercise your rights and contact us: See Sections 16 and 17 of this policy.\n\n"
     "How to complain to the Data Protection Board: See Section 14 of this policy."),

    # Section 3 — Consent (heading keywords: consent, legal basis, withdrawal, section 6)
    (1, "3. Legal Basis for Processing - Consent (Section 6 DPDP Act)"),
    (None,
     "We process your personal data on the following legal bases:\n\n"
     "Consent (Section 6 DPDP Act): When you register, you give free, specific, informed, "
     "and unambiguous consent by affirmatively accepting this Privacy Policy and our Terms "
     "of Service through a clear checkbox action. Consent is not bundled with other terms "
     "in a manner that prevents separate refusal.\n\n"
     "Withdrawal of consent: You may withdraw your consent at any time by deleting your account "
     "via the platform settings or by sending a written withdrawal request to "
     "privacy@globallearn.example. Withdrawal of consent will not affect the lawfulness of "
     "processing carried out before withdrawal.\n\n"
     "Cessation of processing within 30 days: Upon receiving a valid consent withdrawal "
     "request, GlobalLearn will cease all non-essential processing of your personal data within "
     "thirty (30) days. Processing required to fulfil a legal obligation, resolve a dispute, "
     "or enforce our agreements may continue only to the extent strictly necessary.\n\n"
     "What happens to your data after withdrawal: Once non-essential processing ceases, your "
     "personal data will be erased within 90 days of account closure, except where mandatory "
     "legal retention obligations apply. We will send you a confirmation once your data has "
     "been deleted. During the 30-day cessation period, your data will not be used for any "
     "new processing; it will be held in read-only form pending deletion."),

    # Section 4 — Legitimate Uses (heading keywords: legitimate, legal basis, without consent, section 7)
    (1, "4. Legal Basis for Processing - Legitimate Uses Without Consent (Section 7 DPDP Act)"),
    (None,
     "We may process personal data without consent under the following legitimate uses "
     "permitted by Section 7 of the DPDP Act:\n\n"
     "  a) Legal compliance: to comply with a court order, legal obligation, or direction "
     "of a competent authority (e.g., tax or regulatory filings, law enforcement requests).\n\n"
     "  b) Medical or public safety emergency: to respond to a life-threatening situation "
     "affecting a user or another person, to the extent strictly necessary.\n\n"
     "  c) Employment and benefits: where our platform is used in a corporate learning "
     "context, we may process employee personal data to the extent necessary to perform or "
     "terminate employment-related obligations under applicable law.\n\n"
     "  d) Merger or acquisition: if GlobalLearn undergoes a merger, acquisition, or asset "
     "sale, personal data may be transferred to the acquiring entity subject to equivalent "
     "data protection safeguards.\n\n"
     "In each case, we will document the basis for such processing and inform affected "
     "Data Principals where disclosure is not restricted by law."),

    # Section 5 — Consent Manager (heading keywords: consent manager)
    (1, "5. Consent Manager (Section 6(7)-(9) DPDP Act and DPDP Rules 2025 Rule 4)"),
    (None,
     "GlobalLearn supports the use of registered Consent Managers to allow Data Principals "
     "to manage, review, and withdraw their consent centrally, as permitted by Section 6(7)-(9) "
     "of the DPDP Act and DPDP Rules 2025 Rule 4.\n\n"
     "What is a Consent Manager: A Consent Manager is a registered entity that maintains "
     "a user-accessible platform through which Data Principals may give, review, and withdraw "
     "consent with multiple Data Fiduciaries. Consent Managers are registered with and "
     "regulated by the Data Protection Board of India.\n\n"
     "How to use a Consent Manager with GlobalLearn: If you choose to manage your consent "
     "through a registered Consent Manager rather than directly through our platform, you may:\n"
     "  (a) Access your consent records via the Consent Manager's interface.\n"
     "  (b) Modify or withdraw consent for specific processing activities.\n"
     "  (c) Request GlobalLearn to honour instructions received from your Consent Manager.\n\n"
     "Direct management: You may also manage consent directly through your GlobalLearn "
     "account settings at any time without using a Consent Manager. Direct withdrawal and "
     "management through our platform is always available.\n\n"
     "Note: GlobalLearn currently also manages consent directly. Use of a Consent Manager "
     "is optional and does not affect your rights or the protections available to you."),

    # Section 6 — Vendors/Processors (heading keywords: data processor, third-party, processor, section 8 dpdp)
    (1, "6. Third-Party Data Processors (Section 8 DPDP Act)"),
    (None,
     "To deliver our platform, we engage the following categories of third-party Data "
     "Processors. Each processes personal data only under our written instructions and under "
     "a Data Processing Agreement that imposes security, confidentiality, audit rights, "
     "sub-processor restrictions, and breach notification obligations:\n\n"
     "  - Cloud infrastructure provider: hosts the platform and its databases on servers in "
     "India and Singapore. Servers in Singapore process data of international users only.\n\n"
     "  - AI engine provider: powers our personalised learning recommendation and assessment "
     "features. This processor accesses learner behaviour data solely for model inference "
     "and is prohibited from training on or retaining personal data.\n\n"
     "  - Payment gateway: processes subscription billing. We do not store full card numbers; "
     "the gateway handles payment data under PCI-DSS controls.\n\n"
     "  - Transactional email provider: delivers account verification, password reset, and "
     "course completion notifications. Processes only recipient email and message content.\n\n"
     "  - Video streaming provider: hosts recorded course content. Accesses only technical "
     "metadata (user ID and session token) to authenticate playback requests.\n\n"
     "  - Analytics processor: provides aggregated, anonymised platform usage analytics. "
     "Receives pseudonymised event data only; no directly identifying personal data.\n\n"
     "GlobalLearn remains fully responsible as the Data Fiduciary for all processing by "
     "these processors. We review processor compliance periodically and will not engage any "
     "processor that cannot provide adequate data protection guarantees."),

    # Section 7 — Transfer (heading keywords: cross-border, transfer, section 16, storage location)
    (1, "7. Cross-Border Data Transfers and Storage Location (Section 16 DPDP Act and DPDP Rules 2025 Rule 15)"),
    (None,
     "Data storage locations: Personal data of users who are Indian residents is stored on "
     "servers located in India. Personal data of international users (outside India) is "
     "processed and stored on servers located in Singapore, subject to the safeguards "
     "described below.\n\n"
     "Cross-border transfers: We transfer personal data outside India only to the extent "
     "necessary to serve international users and to deliver services through our Singapore-based "
     "cloud infrastructure provider, as described in Section 6.\n\n"
     "Safeguards for cross-border transfers:\n"
     "  (a) We transfer data only to countries and entities approved or not restricted under "
     "Section 16 of the DPDP Act and any Central Government notification.\n"
     "  (b) All cross-border transfers are governed by Data Processing Agreements that impose "
     "security and confidentiality obligations at least equivalent to those in this policy.\n"
     "  (c) We monitor changes to Central Government restrictions and will cease transfers to "
     "any restricted country immediately upon notification.\n\n"
     "We do not transfer personal data of Indian residents outside India except where "
     "expressly required to deliver the service or comply with a legal obligation. We will "
     "update this policy and obtain required approvals before making any material change to "
     "our transfer arrangements."),

    # Section 8 — Security (heading keywords: security safeguard, section 8(5), rule 6)
    (1, "8. Security Safeguards (Section 8(5) DPDP Act and DPDP Rules 2025 Rule 6)"),
    (None,
     "We implement reasonable and appropriate technical and organisational security safeguards "
     "to protect personal data from unauthorised access, disclosure, alteration, or destruction:\n\n"
     "  - Encryption in transit: all data transmitted between your browser or client and our "
     "servers is encrypted using TLS 1.3.\n\n"
     "  - Encryption at rest: personal data stored in our databases and backups is encrypted "
     "using AES-256 or an equivalent industry-standard algorithm.\n\n"
     "  - Access controls: role-based access controls (RBAC) ensure that only authorised "
     "personnel access personal data, based on the principle of least privilege. Access "
     "privileges are reviewed quarterly and revoked promptly on termination.\n\n"
     "  - Monitoring and logging: we maintain audit logs of all access to personal data "
     "systems. Logs are retained for a minimum of one year and monitored continuously.\n\n"
     "  - AI model security: our AI personalisation engine processes learner data in an "
     "isolated environment. Model outputs are not stored alongside raw personal data.\n\n"
     "  - Backup and continuity: we maintain encrypted off-site backups and a tested "
     "business continuity plan to ensure data availability and rapid recovery.\n\n"
     "  - Organisational security measures: all employees with access to personal data "
     "complete mandatory data protection training annually and are bound by confidentiality "
     "obligations. Third-party access is governed by Data Processing Agreements.\n\n"
     "  - Data accuracy: we take reasonable steps to ensure personal data used in decisions "
     "affecting you (including AI recommendations) is accurate and up to date. You may "
     "correct inaccurate data at any time via account settings or by writing to "
     "privacy@globallearn.example."),

    # Section 9 — Retention (heading keywords: retention, erasure, section 8(7), rule 8)
    (1, "9. Retention and Erasure (Section 8(7)-(8) DPDP Act and DPDP Rules 2025 Rule 8)"),
    (None,
     "We retain personal data only for as long as necessary for the stated purposes or as "
     "required by law:\n\n"
     "  - Active accounts: personal data is retained while your account is active and your "
     "subscription is in force.\n\n"
     "  - After account closure: when you delete your account or your subscription ends, we "
     "will erase your personal data within 90 days, unless retention is required by law "
     "(for example, financial records under the Companies Act or GST law).\n\n"
     "  - Inactive accounts: if your account has been inactive (no login or activity) for "
     "24 continuous months, we will send a notice to your registered email address at least "
     "30 days before deleting your data, giving you the opportunity to log in and reactivate "
     "your account to prevent deletion.\n\n"
     "  - Processing and access logs: system audit logs are retained for a minimum of one "
     "year from creation as required under DPDP Rules 2025 Rule 8, after which they are "
     "securely and irreversibly deleted.\n\n"
     "  - AI training data: learner behaviour signals used for AI personalisation are "
     "purged when your account is deleted. We do not retain personal data for AI model "
     "training beyond the period your account is active.\n\n"
     "  - Consent records: records evidencing your consent are retained for the duration of "
     "our relationship and for three years thereafter in case of legal dispute.\n\n"
     "When data is erased, it is deleted from primary databases and purged from backups "
     "within the next scheduled backup cycle (maximum 30 additional days)."),

    # Section 10 — Children (heading keywords: children, persons under 18, parental consent, section 9)
    (1, "10. Children's Data and Parental Consent (Section 9 DPDP Act and DPDP Rules 2025 Rule 10)"),
    (None,
     "GlobalLearn's platform is accessible to learners of all ages, including children under "
     "18. We comply fully with Section 9 of the DPDP Act and DPDP Rules 2025 Rule 10.\n\n"
     "Verifiable parental consent: Before processing personal data of any child under 18, "
     "GlobalLearn obtains verifiable consent from the child's parent or lawful guardian. "
     "Consent is obtained through the following process:\n"
     "  (a) The parent or guardian registers a parent account and provides their full name, "
     "email address, and a declaration of guardianship.\n"
     "  (b) We verify the parent's identity through a one-time password (OTP) sent to their "
     "registered mobile number or email address.\n"
     "  (c) The parent affirmatively consents to the child's account creation and data "
     "processing by confirming acceptance of this policy on behalf of the child.\n"
     "  (d) The child's account is linked to the parent account, allowing the parent to "
     "view, correct, or delete the child's data at any time.\n\n"
     "Age verification: We use a date-of-birth declaration combined with parent account "
     "verification to determine whether a user is under 18. We reserve the right to "
     "request additional verification documents if we have reason to doubt the accuracy "
     "of the declaration.\n\n"
     "Prohibition of tracking and advertising: We do not track the online behaviour of "
     "children across third-party websites or apps. We do not serve targeted or behavioural "
     "advertising to children. We do not process children's data in any manner likely to "
     "cause harm to their well-being, physical safety, or mental health.\n\n"
     "Parental controls: Parents and guardians may, at any time:\n"
     "  (a) Access a summary of all personal data held about their child.\n"
     "  (b) Correct or update the child's personal data.\n"
     "  (c) Request erasure of the child's account and data.\n"
     "  (d) Withdraw consent for all data processing by deleting the child's account.\n\n"
     "Requests for parental controls: Contact privacy@globallearn.example with the subject "
     "line 'Child Account Request' and include the child's username and your registered "
     "parent account email."),

    # Section 11 — Disability (heading keywords: disability, guardian authority, section 9(1), rule 11)
    (1, "11. Data of Persons with Disability - Guardian Authority Verification (Section 9(1) DPDP Act and DPDP Rules 2025 Rule 11)"),
    (None,
     "Where a learner is a person with disability who has a lawful guardian, GlobalLearn "
     "will process their personal data only after obtaining verifiable consent from that "
     "guardian and completing a formal guardian authority verification process, as required "
     "by Section 9(1) of the DPDP Act and DPDP Rules 2025 Rule 11.\n\n"
     "Documents required for guardian authority verification:\n"
     "  (a) Proof of guardian's identity: a valid government-issued photo identity document "
     "(Aadhaar, PAN, passport, or driving licence).\n"
     "  (b) Proof of data principal's identity: a valid identity document for the person "
     "with disability.\n"
     "  (c) Legal authority document: a court order, notarised guardianship deed, power of "
     "attorney, or any other instrument issued by a competent authority establishing the "
     "guardian's legal authority over the data principal.\n"
     "  (d) Disability certificate: a valid certificate issued by a competent medical or "
     "government authority confirming the nature of the disability.\n\n"
     "Verification criteria: We will assess whether (i) the legal instrument is valid, "
     "current, and not revoked; (ii) the identity of the guardian matches the instrument; "
     "and (iii) the scope of authority covers consent to personal data processing.\n\n"
     "Verification procedure:\n"
     "  Step 1 - Request: the guardian emails privacy@globallearn.example with the subject "
     "line 'Guardian Consent Verification', providing the data principal's name, disability "
     "details, and the guardian's contact information.\n"
     "  Step 2 - Document submission: the guardian submits the required documents listed above.\n"
     "  Step 3 - Review: GlobalLearn reviews submissions within 10 business days against "
     "the verification criteria and confirms in writing whether authority is verified.\n"
     "  Step 4 - Processing: personal data is processed under guardian consent only after "
     "written confirmation of successful verification.\n"
     "  Step 5 - Records: verified records are retained securely; the guardian may revoke "
     "authority at any time by written notice to privacy@globallearn.example.\n\n"
     "We will not process personal data of a person with disability under guardian consent "
     "until all verification steps are completed."),

    # Section 12 — Breach (heading keywords: breach notification, data breach, section 8(6), rule 7)
    (1, "12. Data Breach Notification (Section 8(6) DPDP Act and DPDP Rules 2025 Rule 7)"),
    (None,
     "Board Intimation within 72 Hours: GlobalLearn will notify the Data Protection Board "
     "of India within seventy-two (72) hours of becoming aware of any personal data breach, "
     "as required by Section 8(6) of the DPDP Act and DPDP Rules 2025 Rule 7. This "
     "72-hour Board intimation obligation applies regardless of the number of Data Principals "
     "affected or the severity of the breach.\n\n"
     "In the event of a personal data breach, GlobalLearn will take the following steps:\n\n"
     "  a) Notify affected Data Principals without undue delay: we will inform each affected "
     "individual of the nature and description of the breach, the categories and approximate "
     "volume of personal data affected, the likely consequences of the breach, the measures "
     "taken or proposed to address and mitigate the breach, and a contact point for further "
     "information and assistance.\n\n"
     "  b) Notify the Data Protection Board of India within 72 hours: within seventy-two "
     "(72) hours of becoming aware of a breach, we will submit a formal breach notification "
     "to the Data Protection Board in the prescribed format under DPDP Rules 2025 Rule 7, "
     "including a description of the breach, its likely consequences, the data and persons "
     "affected, and remedial measures taken or planned. If all required details are not "
     "available within 72 hours, we will submit an initial notification within that window "
     "and provide a complete follow-up report as soon as practicable.\n\n"
     "  c) Contain and investigate: we will immediately isolate affected systems, conduct "
     "a root-cause investigation, and implement remedial measures to prevent recurrence.\n\n"
     "  d) Maintain breach records: we maintain an internal register of all personal data "
     "breaches, including those that did not require external notification, for audit and "
     "accountability purposes."),

    # Section 13 — Grievance (heading keywords: grievance, section 8(10), section 13, rule 14(3))
    (1, "13. Grievance Redressal (Section 8(10) and Section 13 DPDP Act and DPDP Rules 2025 Rule 14(3))"),
    (None,
     "If you have any complaint or concern about how we handle your personal data, please "
     "contact our Grievance Officer:\n\n"
     "  Name:         Arjun Mehta\n"
     "  Designation:  Grievance Officer and Data Protection Officer\n"
     "  Organisation: GlobalLearn India Pvt Ltd\n"
     "  Email:        privacy@globallearn.example\n"
     "  Website:      https://globallearn.example/privacy\n\n"
     "We will acknowledge your grievance within 7 days and endeavour to resolve it within "
     "ninety (90) days of receipt. If your grievance is not resolved to your satisfaction "
     "within ninety days, you may escalate your complaint to the Data Protection Board of "
     "India (see Section 14). Data Principals are required to exhaust our grievance "
     "redressal process before approaching the Board, unless the Board directs otherwise."),

    # Section 14 — Board complaint (heading keywords: complaint to, data protection board, board of india)
    (1, "14. Complaint to the Data Protection Board of India (Section 5(1)(iii) DPDP Act)"),
    (None,
     "If your grievance with GlobalLearn is not resolved within ninety days, or if you "
     "believe we have violated the Digital Personal Data Protection Act, 2023, you have the "
     "right to file a complaint with the Data Protection Board of India.\n\n"
     "The Data Protection Board of India is the independent adjudicatory authority established "
     "under Chapter VI of the DPDP Act. Complaints may be filed through the Board's online "
     "portal or in such other manner as may be prescribed. For current guidance on filing a "
     "complaint, refer to the Board's official website or contact the Ministry of Electronics "
     "and Information Technology (MeitY) at www.meity.gov.in. Filing a complaint with the "
     "Board is free of charge. Children and persons with disability may file complaints through "
     "their parent or guardian."),

    # Section 15 — SDF (heading keywords: significant data fiduciary, section 10, sdf status)
    (1, "15. Significant Data Fiduciary Status (Section 10 DPDP Act and DPDP Rules 2025 Rule 13)"),
    (None,
     "GlobalLearn India Pvt Ltd has been notified by the Central Government of India as a "
     "Significant Data Fiduciary (SDF) under Section 10 of the Digital Personal Data "
     "Protection Act, 2023. Accordingly, the following additional obligations under "
     "Section 10 and DPDP Rules 2025 Rule 13 apply to GlobalLearn:\n\n"
     "  (a) India-resident Data Protection Officer (DPO): GlobalLearn has appointed "
     "Arjun Mehta as its resident DPO. The DPO is a key managerial person responsible for "
     "data protection compliance and for acting as the point of contact for the Data "
     "Protection Board. The DPO's contact details are published in Section 17 of this policy.\n\n"
     "  (b) Data Protection Impact Assessment (DPIA): GlobalLearn conducts periodic DPIAs "
     "before implementing new AI features, new data collection categories, or significant "
     "changes to processing activities that may affect Data Principal rights.\n\n"
     "  (c) Independent audit: GlobalLearn undergoes an annual independent data protection "
     "audit by a qualified third-party auditor. Audit findings are reported to the DPO and "
     "material issues are reported to the Data Protection Board as required.\n\n"
     "  (d) Algorithmic accountability: our AI personalisation and assessment algorithms "
     "are reviewed periodically to ensure they do not pose risk to the rights and well-being "
     "of Data Principals, including children. We do not use fully automated decision-making "
     "that has significant effects on a learner without human review.\n\n"
     "  (e) Restrictions on sensitive data transfers: as required under Rule 13, GlobalLearn "
     "does not transfer any notified categories of sensitive personal data outside India "
     "without Central Government approval."),

    # Section 16 — Rights (heading keywords: your rights, rights as a data principal, section 11, section 14)
    (1, "16. Your Rights as a Data Principal (Sections 11-14 DPDP Act and DPDP Rules 2025 Rule 14)"),
    (None,
     "You have the following rights regarding your personal data. To exercise any right, "
     "write to privacy@globallearn.example with the subject line 'Data Rights Request' "
     "and include your name and registered email address:\n\n"
     "  a) Right to access: request a summary of the personal data we hold about you and "
     "a list of Data Processors and third parties with whom it has been shared.\n\n"
     "  b) Right to correction: request correction, completion, or updating of any inaccurate, "
     "incomplete, or outdated personal data we hold about you.\n\n"
     "  c) Right to erasure: request deletion of your personal data where it is no longer "
     "necessary for the purposes for which it was collected, subject to legal retention "
     "requirements.\n\n"
     "  d) Right to withdraw consent: withdraw consent at any time by deleting your account "
     "or sending a written withdrawal request to privacy@globallearn.example. Withdrawal "
     "does not affect the lawfulness of processing carried out before withdrawal.\n\n"
     "  e) Right to information about processors: request details of the Data Processors to "
     "whom your personal data has been disclosed and for what purpose.\n\n"
     "  f) Right of nomination: nominate another individual to exercise your data rights on "
     "your behalf in the event of your death or incapacity, by submitting a written nomination "
     "to privacy@globallearn.example.\n\n"
     "  g) Rights for children and guardians: parents and guardians may exercise all of the "
     "above rights on behalf of a child under 18. See Section 10 for the parental controls "
     "process.\n\n"
     "We will acknowledge your request within 7 days and respond fully within 30 days. "
     "We do not charge a fee for reasonable requests. We may ask you to verify your identity "
     "before processing your request."),

    # Section 17 — Contact / DPO (heading keywords: designated contact, section 8(9), rule 9, dpo)
    (1, "17. Designated Contact for Personal Data Processing Questions (Section 8(9) DPDP Act and DPDP Rules 2025 Rule 9)"),
    (None,
     "As required by Section 8(9) of the DPDP Act and Rule 9 of the DPDP Rules 2025, "
     "GlobalLearn has designated the following person — who also serves as our DPO under "
     "our Significant Data Fiduciary obligations — to answer any questions about the "
     "processing of personal data. This contact is prominently published on our website, "
     "application, and in this policy:\n\n"
     "  Designated Contact and Data Protection Officer (DPO):\n"
     "  Name:         Arjun Mehta\n"
     "  Designation:  Data Protection Officer / Grievance Officer\n"
     "  Organisation: GlobalLearn India Pvt Ltd\n"
     "  Email:        privacy@globallearn.example\n"
     "  Website:      https://globallearn.example/privacy\n"
     "  Available:    Monday to Friday, 9:00 AM to 6:00 PM IST\n\n"
     "You may contact this person for any question relating to: what personal data we "
     "collect; why and how we process it; who we share it with; how long we keep it; how "
     "to exercise your rights; AI personalisation; children's data; or any other matter "
     "concerning personal data processing by GlobalLearn."),

    # Section 18 — Updates
    (1, "18. Updates to This Policy"),
    (None,
     "This policy is reviewed at least annually and whenever there is a material change "
     "to our processing activities. Material changes will be communicated to registered "
     "users by email at least 30 days before they take effect. Parents and guardians of "
     "children under 18 will be separately notified of any change that affects children's "
     "data processing. Continued use of the platform after the effective date of a revised "
     "policy constitutes acceptance of the revised terms. The current version is always "
     "available at https://globallearn.example/privacy."),
]


def main() -> None:
    doc = Document()

    heading = doc.add_heading(TITLE, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(META)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.size = Pt(9)
        run.font.italic = True

    for level, text in BODY:
        if level == 1:
            doc.add_heading(text, level=1)
        else:
            doc.add_paragraph(text)

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print()
    print("=== QUESTIONNAIRE ANSWERS TO USE ===")
    print("Company name:      GlobalLearn India Pvt Ltd")
    print("Website:           https://globallearn.example")
    print("Email:             privacy@globallearn.example")
    print("Grievance officer: Arjun Mehta")
    print("Processing type:   Website / app with user accounts")
    print("Audience:          General public")
    print("Data collected:    Account and login details")
    print("Uses AI:           Yes  <-- critical for CONSENT_MANAGER rule")
    print("Data storage:      Multiple countries  <-- critical for TRANSFER rule")
    print("Third parties:     Yes")
    print("Analytics:         Yes")
    print("Retention:         90 days after account closure")
    print("Under 18:          Yes  <-- critical for CHILD + DISABILITY rules")
    print("Intl users:        Yes")
    print("Security:          Yes")
    print()
    print("Expected result:   COMPLIANT — all 16 rules applicable and satisfied")


if __name__ == "__main__":
    main()
